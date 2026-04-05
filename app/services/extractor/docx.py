from pathlib import Path
from typing import Iterator

from app.core.logging import get_logger
from app.services.extractor.base import BaseExtractor, ExtractedPage

log = get_logger(__name__)


class DocxExtractor(BaseExtractor):
    """Extract text from .docx files, preserving paragraph order."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def extract(self, file_path: Path) -> Iterator[ExtractedPage]:
        log.debug("docx.start", file=str(file_path))
        try:
            from docx import Document  # lazy import

            doc = Document(str(file_path))
        except Exception as exc:
            log.error("docx.open_failed", file=str(file_path), error=str(exc))
            raise

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also collect text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        full_text = "\n".join(paragraphs)
        if full_text.strip():
            yield ExtractedPage(page_number=0, text=full_text)
        else:
            log.warning("docx.no_content", file=str(file_path))

        log.debug("docx.done", file=str(file_path), paragraphs=len(paragraphs))
