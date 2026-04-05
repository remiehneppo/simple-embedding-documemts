"""
Unit tests for app.services.extractor_factory and individual extractors.

- factory dispatch by file extension
- PlainTextExtractor: .txt, .md, .csv
- DocxExtractor: .docx
- PdfTextExtractor: text-layer PDF
- Unsupported extensions raise UnsupportedFileTypeError
- PdfOcrExtractor routing (patched has_text_layer — no actual OCR run)
"""

from pathlib import Path
from unittest.mock import patch

import pytest

from app.core.exceptions import UnsupportedFileTypeError
from app.services.extractor.base import ExtractedPage
from app.services.extractor.docx import DocxExtractor
from app.services.extractor.plaintext import PlainTextExtractor
from app.services.extractor.pdf_text import PdfTextExtractor
from app.services.extractor_factory import SUPPORTED_EXTENSIONS, get_extractor


# ── Factory dispatch ─────────────────────────────────────────────────────────

class TestExtractorFactoryDispatch:
    def test_txt_returns_plaintext_extractor(self, tmp_path):
        p = tmp_path / "file.txt"
        p.write_text("hello")
        assert isinstance(get_extractor(p), PlainTextExtractor)

    def test_md_returns_plaintext_extractor(self, tmp_path):
        p = tmp_path / "file.md"
        p.write_text("# Hello")
        assert isinstance(get_extractor(p), PlainTextExtractor)

    def test_csv_returns_plaintext_extractor(self, tmp_path):
        p = tmp_path / "file.csv"
        p.write_text("a,b,c")
        assert isinstance(get_extractor(p), PlainTextExtractor)

    def test_docx_returns_docx_extractor(self, sample_docx):
        assert isinstance(get_extractor(sample_docx), DocxExtractor)

    def test_pdf_with_text_layer_returns_pdf_text_extractor(self, sample_pdf_text):
        extractor = get_extractor(sample_pdf_text)
        assert isinstance(extractor, PdfTextExtractor)

    def test_pdf_without_text_layer_returns_pdf_ocr_extractor(self, tmp_path):
        """Patch has_text_layer to return False to avoid running actual OCR."""
        from app.services.extractor.pdf_ocr import PdfOcrExtractor

        # Create a minimal (but parseable) PDF
        import fitz
        doc = fitz.open()
        doc.new_page()  # blank page → no text layer
        p = tmp_path / "blank.pdf"
        doc.save(str(p))
        doc.close()

        with patch("app.services.extractor_factory.has_text_layer", return_value=False):
            extractor = get_extractor(p)
        assert isinstance(extractor, PdfOcrExtractor)

    def test_unsupported_extension_raises(self, tmp_path):
        p = tmp_path / "file.xyz"
        p.write_text("content")
        with pytest.raises(UnsupportedFileTypeError) as exc_info:
            get_extractor(p)
        assert ".xyz" in str(exc_info.value)

    def test_unsupported_extension_exe_raises(self, tmp_path):
        p = tmp_path / "malware.exe"
        p.write_text("content")
        with pytest.raises(UnsupportedFileTypeError):
            get_extractor(p)

    def test_supported_extensions_set(self):
        assert ".pdf" in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".txt" in SUPPORTED_EXTENSIONS
        assert ".md" in SUPPORTED_EXTENSIONS
        assert ".csv" in SUPPORTED_EXTENSIONS


# ── PlainTextExtractor ────────────────────────────────────────────────────────

class TestPlainTextExtractor:
    def test_yields_single_page(self, sample_txt):
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(sample_txt))
        assert len(pages) == 1

    def test_page_number_is_zero(self, sample_txt):
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(sample_txt))
        assert pages[0].page_number == 0

    def test_text_contains_content(self, sample_txt):
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(sample_txt))
        assert "quick brown fox" in pages[0].text

    def test_md_file_extracted(self, sample_md):
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(sample_md))
        assert len(pages) == 1
        assert "Heading" in pages[0].text

    def test_empty_file_yields_nothing(self, tmp_path):
        p = tmp_path / "empty.txt"
        p.write_text("")
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(p))
        assert pages == []

    def test_whitespace_only_file_yields_nothing(self, tmp_path):
        p = tmp_path / "spaces.txt"
        p.write_text("   \n\t  \n")
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(p))
        assert pages == []

    def test_unicode_content(self, tmp_path):
        p = tmp_path / "unicode.txt"
        p.write_text("Xin chào thế giới. Привет мир.", encoding="utf-8")
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(p))
        assert len(pages) == 1
        assert "Xin chào" in pages[0].text

    def test_returns_extracted_page_type(self, sample_txt):
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(sample_txt))
        assert isinstance(pages[0], ExtractedPage)

    def test_csv_extracted_as_text(self, tmp_path):
        p = tmp_path / "data.csv"
        p.write_text("name,age\nAlice,30\nBob,25")
        extractor = PlainTextExtractor()
        pages = list(extractor.extract(p))
        assert "Alice" in pages[0].text


# ── DocxExtractor ─────────────────────────────────────────────────────────────

class TestDocxExtractor:
    def test_yields_single_page(self, sample_docx):
        extractor = DocxExtractor()
        pages = list(extractor.extract(sample_docx))
        assert len(pages) == 1

    def test_page_number_is_zero(self, sample_docx):
        extractor = DocxExtractor()
        pages = list(extractor.extract(sample_docx))
        assert pages[0].page_number == 0

    def test_all_paragraphs_extracted(self, sample_docx):
        extractor = DocxExtractor()
        pages = list(extractor.extract(sample_docx))
        text = pages[0].text
        assert "First paragraph" in text
        assert "Second paragraph" in text
        assert "Third paragraph" in text

    def test_docx_with_table(self, tmp_path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Before table.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell A1"
        table.cell(0, 1).text = "Cell B1"
        table.cell(1, 0).text = "Cell A2"
        table.cell(1, 1).text = "Cell B2"
        p = tmp_path / "table.docx"
        doc.save(str(p))

        extractor = DocxExtractor()
        pages = list(extractor.extract(p))
        assert "Cell A1" in pages[0].text
        assert "Cell B2" in pages[0].text

    def test_empty_docx_yields_nothing(self, tmp_path):
        from docx import Document

        doc = Document()
        p = tmp_path / "empty.docx"
        doc.save(str(p))

        extractor = DocxExtractor()
        pages = list(extractor.extract(p))
        assert pages == []


# ── PdfTextExtractor ─────────────────────────────────────────────────────────

class TestPdfTextExtractor:
    def test_yields_one_page(self, sample_pdf_text):
        extractor = PdfTextExtractor()
        pages = list(extractor.extract(sample_pdf_text))
        assert len(pages) >= 1

    def test_page_number_starts_at_one(self, sample_pdf_text):
        extractor = PdfTextExtractor()
        pages = list(extractor.extract(sample_pdf_text))
        assert pages[0].page_number == 1

    def test_text_extracted(self, sample_pdf_text):
        extractor = PdfTextExtractor()
        pages = list(extractor.extract(sample_pdf_text))
        combined = " ".join(p.text for p in pages)
        assert "text layer" in combined.lower()

    def test_blank_page_skipped(self, tmp_path):
        import fitz

        doc = fitz.open()
        doc.new_page()  # blank
        page2 = doc.new_page()
        page2.insert_text((50, 72), "Second page content.", fontsize=12)
        p = tmp_path / "mixed.pdf"
        doc.save(str(p))
        doc.close()

        extractor = PdfTextExtractor()
        pages = list(extractor.extract(p))
        # Only the non-blank page should be yielded
        assert len(pages) == 1
        assert pages[0].page_number == 2

    def test_multi_page_pdf(self, tmp_path):
        import fitz

        doc = fitz.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((50, 72), f"Page {i + 1} content.", fontsize=12)
        p = tmp_path / "multi.pdf"
        doc.save(str(p))
        doc.close()

        extractor = PdfTextExtractor()
        pages = list(extractor.extract(p))
        assert len(pages) == 3
        page_numbers = [p.page_number for p in pages]
        assert page_numbers == [1, 2, 3]


# ── has_text_layer ────────────────────────────────────────────────────────────

class TestHasTextLayer:
    def test_returns_true_for_text_pdf(self, sample_pdf_text):
        from app.services.extractor.pdf_text import has_text_layer

        assert has_text_layer(sample_pdf_text) is True

    def test_returns_false_for_blank_pdf(self, tmp_path):
        import fitz
        from app.services.extractor.pdf_text import has_text_layer

        doc = fitz.open()
        doc.new_page()
        p = tmp_path / "blank.pdf"
        doc.save(str(p))
        doc.close()

        assert has_text_layer(p) is False

    def test_returns_false_for_nonexistent_file(self, tmp_path):
        from app.services.extractor.pdf_text import has_text_layer

        assert has_text_layer(tmp_path / "ghost.pdf") is False
