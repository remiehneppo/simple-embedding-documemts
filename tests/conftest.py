"""
Shared pytest fixtures for the simple-embed test suite.

Isolation strategy
------------------
- Every test gets a fresh ChromaDB collection stored in a pytest tmp_path.
- `settings` attributes (chroma_path, documents_path) are monkeypatched per test.
- The ChromaDB global singletons (_client, _collection) are reset before and
  after every test via the autouse `reset_chroma_state` fixture.
- The Ollama embedder is NOT called for unit / service tests — a deterministic
  fake is injected via `monkeypatch`.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import AsyncIterator
from unittest.mock import AsyncMock, patch

import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient

# ── Helpers ──────────────────────────────────────────────────────────────────

FAKE_VECTOR_DIM = 768


def make_fake_embeddings(texts: list[str]) -> list[list[float]]:
    """Return deterministic unit vectors (one per text, cycling through values)."""
    result = []
    for i, _ in enumerate(texts):
        # Build a simple vector where position i % dim = 1.0, rest = 0.0
        vec = [0.0] * FAKE_VECTOR_DIM
        vec[i % FAKE_VECTOR_DIM] = 1.0
        result.append(vec)
    return result


async def _async_fake_embed(texts: list[str]) -> list[list[float]]:
    return make_fake_embeddings(texts)


# ── ChromaDB isolation ───────────────────────────────────────────────────────

@pytest.fixture(autouse=True)
def reset_chroma_state():
    """Reset ChromaDB module-level singletons before and after every test."""
    from app.db.chroma import reset_collection_cache

    reset_collection_cache()
    yield
    reset_collection_cache()


@pytest.fixture()
def chroma_tmp(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    """
    Redirect ChromaDB and document storage to a per-test temp directory.
    Returns the tmp_path root.
    """
    from app.core import config

    monkeypatch.setattr(config.settings, "chroma_path", tmp_path / "chroma_db")
    monkeypatch.setattr(config.settings, "documents_path", tmp_path / "documents")
    (tmp_path / "documents").mkdir(parents=True, exist_ok=True)
    return tmp_path


# ── Embedding mock ───────────────────────────────────────────────────────────

@pytest.fixture()
def patch_embedder(monkeypatch: pytest.MonkeyPatch):
    """
    Patch embed_texts in every module that imports it so no Ollama call is made.
    Returns the AsyncMock so tests can inspect calls.
    """
    mock = AsyncMock(side_effect=_async_fake_embed)
    monkeypatch.setattr("app.services.embedder.embed_texts", mock)
    monkeypatch.setattr("app.services.document_service.embed_texts", mock)
    monkeypatch.setattr("app.services.search_service.embed_texts", mock)
    return mock


# ── Sample documents ─────────────────────────────────────────────────────────

@pytest.fixture()
def sample_txt(tmp_path: Path) -> Path:
    """A plain-text file with several sentences."""
    p = tmp_path / "sample.txt"
    p.write_text(
        "The quick brown fox jumps over the lazy dog. "
        "Pack my box with five dozen liquor jugs. "
        "How vexingly quick daft zebras jump. "
        "The five boxing wizards jump quickly.",
        encoding="utf-8",
    )
    return p


@pytest.fixture()
def sample_md(tmp_path: Path) -> Path:
    p = tmp_path / "sample.md"
    p.write_text(
        "# Heading\n\nThis is a markdown document.\n\nIt has multiple paragraphs.",
        encoding="utf-8",
    )
    return p


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """A minimal DOCX file with a few paragraphs."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph of the DOCX document.")
    doc.add_paragraph("Second paragraph with more content for testing purposes.")
    doc.add_paragraph("Third paragraph to ensure multiple chunks are possible.")
    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def sample_pdf_text(tmp_path: Path) -> Path:
    """A PDF with a native text layer (PyMuPDF)."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (50, 72),
        "This is a PDF document with a text layer. "
        "It contains enough words to be chunked properly. "
        "Semantic search will be able to find relevant passages.",
        fontsize=12,
    )
    p = tmp_path / "sample_text.pdf"
    doc.save(str(p))
    doc.close()
    return p


# ── FastAPI test client ───────────────────────────────────────────────────────

@pytest_asyncio.fixture()
async def api_client(
    chroma_tmp: Path,
    patch_embedder,
) -> AsyncIterator[AsyncClient]:
    """
    Async HTTPX client wired to the FastAPI app.
    ChromaDB is isolated (tmp dir) and the embedder is mocked.
    """
    from app.main import app

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://testserver"
    ) as client:
        yield client
